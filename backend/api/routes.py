from fastapi import APIRouter, File, UploadFile, Form, HTTPException
from typing import Optional
from pydantic import BaseModel
import io
import base64
import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from services.docling_parser import parse_pdf_from_buffer
from services.langgraph_agent import run_agent
from services.linkedin_scraper import fetch_and_condense_job

router = APIRouter()

class ScrapeRequest(BaseModel):
    url: str

@router.post("/scrape-job")
async def scrape_job_endpoint(payload: ScrapeRequest):
    """
    Endpoint to scrape a job description from a LinkedIn or general job URL
    and return the condensed requirements/responsibilities.
    """
    url = payload.url.strip()
    if not url:
        raise HTTPException(status_code=400, detail="URL cannot be empty")
    try:
        scraped_data = await fetch_and_condense_job(url)
        return scraped_data
    except Exception as e:
        import traceback
        error_detail = traceback.format_exc()
        raise HTTPException(status_code=500, detail=f"Failed to scrape and condense job details: {str(e)}\n{error_detail}")


def create_cover_letter_docx(cl_data: dict, company_name: str, hm_name: str) -> str:
    from docx.shared import Pt, Inches, RGBColor
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn
    import re
    
    doc = Document()
    
    # Optimize fonts and margins for 1-page fit
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    
    def set_bottom_border(paragraph):
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12') # thickness
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_markdown_paragraph(doc, text):
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                p.add_run(part[2:-2]).bold = True
            else:
                p.add_run(part)
        return p

    # Candidate Name (Header)
    name = cl_data.get("candidate_name", "Candidate Name")
    p_name = doc.add_paragraph()
    r_name = p_name.add_run(name)
    r_name.bold = True
    r_name.font.size = Pt(18)
    r_name.font.color.rgb = RGBColor(30, 41, 59) # Dark slate header
    p_name.paragraph_format.space_after = Pt(2)
    
    title = cl_data.get("candidate_title", "")
    if title:
        p_title = doc.add_paragraph()
        r_title = p_title.add_run(title)
        r_title.italic = True
        r_title.font.size = Pt(11)
        r_title.font.color.rgb = RGBColor(100, 116, 139) # Slate subtitle
        p_title.paragraph_format.space_after = Pt(4)
        
    p_contact = doc.add_paragraph(cl_data.get("contact_info", ""))
    p_contact.paragraph_format.space_after = Pt(12)
    set_bottom_border(p_contact)
    
    # Date
    day = datetime.datetime.today().day
    suffix = 'th' if 11 <= day <= 13 else {1:'st',2:'nd',3:'rd'}.get(day % 10, 'th')
    formatted_date = datetime.datetime.today().strftime(f'%d{suffix} %B %Y')
    
    p_date = doc.add_paragraph(formatted_date)
    p_date.paragraph_format.space_before = Pt(8)
    p_date.paragraph_format.space_after = Pt(12)
    
    # Recipient Block - Handle HM name and Company cleanly without duplicate lines
    clean_hm = (hm_name or "").strip()
    clean_company = (company_name or "").strip()
    
    if clean_hm and clean_hm.lower() != "hiring manager":
        p_hm = doc.add_paragraph()
        r_hm = p_hm.add_run(clean_hm)
        r_hm.bold = True
        p_hm.paragraph_format.space_after = Pt(0)
        
        if clean_company:
            p_comp = doc.add_paragraph(f"Hiring Manager, {clean_company}")
            p_comp.paragraph_format.space_after = Pt(14)
        else:
            p_hm.paragraph_format.space_after = Pt(14)
    else:
        p_hm = doc.add_paragraph()
        r_hm = p_hm.add_run("Hiring Manager")
        r_hm.bold = True
        if clean_company:
            p_hm.paragraph_format.space_after = Pt(0)
            p_comp = doc.add_paragraph(clean_company)
            p_comp.paragraph_format.space_after = Pt(14)
        else:
            p_hm.paragraph_format.space_after = Pt(14)
    
    # Greeting
    p_greet = doc.add_paragraph(cl_data.get("greeting", "Dear Hiring Team,"))
    p_greet.paragraph_format.space_after = Pt(10)
    
    # Body Paragraphs
    for p_text in cl_data.get("body_paragraphs", []):
        p_body = add_markdown_paragraph(doc, p_text)
        p_body.paragraph_format.space_after = Pt(10)
        p_body.paragraph_format.line_spacing = 1.15
        
    # Sign-off Block
    sign_off_raw = cl_data.get("sign_off", f"Sincerely,\n{name}")
    if "\n" in sign_off_raw:
        parts = sign_off_raw.split("\n", 1)
        valediction, closing_name = parts[0].strip(), parts[1].strip()
    else:
        valediction, closing_name = sign_off_raw, name

    p_val = doc.add_paragraph(valediction)
    p_val.paragraph_format.space_before = Pt(12)
    p_val.paragraph_format.space_after = Pt(24) # Vertical gap for signature
    
    p_close = doc.add_paragraph()
    r_close = p_close.add_run(closing_name)
    r_close.bold = True
    p_close.paragraph_format.space_after = Pt(0)
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return base64.b64encode(buffer.read()).decode('utf-8')

@router.post("/upload")
async def upload_resume(
    resume: UploadFile = File(...), 
    jd: Optional[str] = Form(None),
    company_name: Optional[str] = Form(None),
    hiring_manager: Optional[str] = Form(None),
    job_url: Optional[str] = Form(None)
):
    """
    Endpoint to process an uploaded PDF resume, Job Description (or Job URL), Company Name, and HM.
    """
    if not resume.filename.endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are supported")
    
    try:
        content = await resume.read()
        parsed_markdown = parse_pdf_from_buffer(content)
        
        resolved_jd = jd
        resolved_company = company_name or ""
        
        if (not resolved_jd or not resolved_jd.strip()) and job_url:
            scraped = await fetch_and_condense_job(job_url)
            resolved_jd = scraped.get("condensed_jd", "")
            # Do not auto-populate company name from scraped details to avoid cover letter triggers
            # if not resolved_company and scraped.get("company"):
            #     resolved_company = scraped.get("company")
                
        if not resolved_jd or not resolved_jd.strip():
            raise HTTPException(status_code=400, detail="Either Job Description text or a valid Job URL must be provided")
            
        agent_result = await run_agent(parsed_markdown, resolved_jd, resolved_company, hiring_manager or "")

        
        if agent_result.get("status") == "failed":
            raise Exception(f"Agent failed: {agent_result.get('error')}")
            
        cover_letter_docx = None
        cl_data = agent_result.get("cover_letter_data", {})
        if cl_data and not cl_data.get("error"):
            cover_letter_docx = create_cover_letter_docx(cl_data, resolved_company, hiring_manager or "")
            
        return {
            "status": "success",
            "filename": resume.filename,
            "parsed_content": parsed_markdown,
            "suggestions": agent_result.get("suggestions", []),
            "ats_score": agent_result.get("ats_score", 0),
            "projected_score": agent_result.get("projected_score", 0),
            "missing_keywords": agent_result.get("missing_keywords", []),
            "match_tier": agent_result.get("match_tier", "Unknown"),
            "cover_letter_docx": cover_letter_docx,
            "cover_letter_data": cl_data if (cl_data and not cl_data.get("error")) else None
        }
    except Exception as e:
        import traceback
        error_detail = traceback.format_exc()
        raise HTTPException(status_code=500, detail=f"Error processing document: {str(e)}\n\nTraceback:\n{error_detail}")
